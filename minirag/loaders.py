"""Document loaders for different file types."""

import json
import os
from pathlib import Path
from typing import List, Union
from urllib.parse import urlparse

from haystack import Document

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    requests = None
    BeautifulSoup = None


class DocumentLoader:
    """Load documents from various file types and sources."""

    def load_documents(
        self, source: Union[str, Path], source_type: str = "auto"
    ) -> List[Document]:
        """Load documents from a source with automatic type detection."""
        if source_type == "auto":
            source_type = self._detect_type(source)

        if source_type == "jsonl":
            return self._load_jsonl(source)
        elif source_type == "pdf":
            return self._load_pdf(source)
        elif source_type == "docx":
            return self._load_docx(source)
        elif source_type == "txt":
            return self._load_txt(source)
        elif source_type == "url":
            return self._load_url(source)
        elif source_type == "directory":
            return self._load_directory(source)
        else:
            raise ValueError(f"Unsupported source type: {source_type}")

    def _detect_type(self, source: Union[str, Path]) -> str:
        """Auto-detect the source type."""
        source_str = str(source)

        # Check if it's a URL
        if source_str.startswith(("http://", "https://")):
            return "url"

        # Check if it's a directory
        if os.path.isdir(source_str):
            return "directory"

        # Check file extension
        if source_str.endswith(".jsonl"):
            return "jsonl"
        elif source_str.endswith(".pdf"):
            return "pdf"
        elif source_str.endswith(".docx"):
            return "docx"
        elif source_str.endswith(".txt"):
            return "txt"
        else:
            raise ValueError(f"Cannot auto-detect type for: {source}")

    def _load_jsonl(self, file_path: Union[str, Path]) -> List[Document]:
        """Load documents from JSONL file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"JSONL file not found: {file_path}")

        docs = []
        with open(file_path, "r", encoding="utf-8") as f:
            for line_num, line in enumerate(f, 1):
                if not line.strip():
                    continue
                try:
                    data = json.loads(line)
                    if "id" not in data or "content" not in data:
                        raise ValueError(
                            f"Missing 'id' or 'content' in line {line_num}"
                        )
                    docs.append(Document(id=data["id"], content=data["content"]))
                except json.JSONDecodeError as e:
                    print(f"Warning: Invalid JSON in line {line_num}: {e}")
                    continue

        if not docs:
            raise ValueError(f"No valid documents found in {file_path}")
        return docs

    def _load_pdf(self, file_path: Union[str, Path]) -> List[Document]:
        """Load document from PDF file."""
        if not pypdf:
            raise ImportError(
                "pypdf is required for PDF loading. Install with: pip install pypdf"
            )

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")

        docs = []
        with open(file_path, "rb") as file:
            reader = pypdf.PdfReader(file)
            content_parts = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    content_parts.append(text.strip())

            if content_parts:
                content = "\n\n".join(content_parts)
                file_name = Path(file_path).stem
                docs.append(Document(id=f"pdf_{file_name}", content=content))

        if not docs:
            raise ValueError(f"No text content found in PDF: {file_path}")
        return docs

    def _load_docx(self, file_path: Union[str, Path]) -> List[Document]:
        """Load document from Word DOCX file."""
        if not DocxDocument:
            raise ImportError(
                "python-docx is required for DOCX loading. Install with: pip install python-docx"
            )

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        if paragraphs:
            content = "\n\n".join(paragraphs)
            file_name = Path(file_path).stem
            return [Document(id=f"docx_{file_name}", content=content)]
        else:
            raise ValueError(f"No text content found in DOCX: {file_path}")

    def _load_txt(self, file_path: Union[str, Path]) -> List[Document]:
        """Load document from text file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Text file not found: {file_path}")

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().strip()

        if content:
            file_name = Path(file_path).stem
            return [Document(id=f"txt_{file_name}", content=content)]
        else:
            raise ValueError(f"No content found in text file: {file_path}")

    def _load_url(self, url: str) -> List[Document]:
        """Load document from web URL."""
        if not requests or not BeautifulSoup:
            raise ImportError(
                "requests and beautifulsoup4 are required for URL loading. Install with: pip install requests beautifulsoup4"
            )

        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
        except requests.RequestException as e:
            raise ValueError(f"Failed to fetch URL {url}: {e}")

        # Parse HTML content
        soup = BeautifulSoup(response.text, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Extract text content
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        content = "\n".join(line for line in lines if line)

        if content:
            # Generate ID from URL
            parsed_url = urlparse(url)
            url_id = f"url_{parsed_url.netloc}_{parsed_url.path.replace('/', '_')}"
            return [Document(id=url_id, content=content)]
        else:
            raise ValueError(f"No content extracted from URL: {url}")

    def _load_directory(self, dir_path: Union[str, Path]) -> List[Document]:
        """Load documents from all supported files in a directory."""
        if not os.path.isdir(dir_path):
            raise ValueError(f"Directory not found: {dir_path}")

        all_docs = []
        supported_extensions = {".jsonl", ".pdf", ".docx", ".txt"}

        for file_path in Path(dir_path).rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    docs = self.load_documents(file_path)
                    all_docs.extend(docs)
                    print(f"Loaded {len(docs)} documents from {file_path}")
                except Exception as e:
                    print(f"Warning: Failed to load {file_path}: {e}")
                    continue

        if not all_docs:
            raise ValueError(f"No supported documents found in directory: {dir_path}")

        return all_docs
